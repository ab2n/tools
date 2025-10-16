#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
scan_arborescence.py
Scan récursif d'un dossier racine, export JSON/Excel/Word.

Usage:
  python scan_arborescence.py            -> ouvrira une boîte de dialogue pour choisir le dossier
  python scan_arborescence.py "D:\BU\ProjetX"  -> utilisera ce dossier comme racine
"""

import os
import sys
import json
from datetime import datetime
from pathlib import Path
import argparse

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    from docx import Document
except ImportError:
    Document = None

from tqdm import tqdm

# ---------------------------
# Helpers
# ---------------------------
def format_timestamp(ts):
    if ts is None:
        return None
    try:
        return datetime.fromtimestamp(ts).isoformat(sep=' ', timespec='seconds')
    except Exception:
        return str(ts)

def safe_relpath(path, start):
    try:
        return os.path.relpath(path, start)
    except Exception:
        return path

def scan_folder(root_path, follow_symlinks=False, show_progress=True):
    """
    Parcourt root_path et renvoie une liste de dicts décrivant chaque fichier.
    """
    root_path = os.path.abspath(root_path)
    results = []

    # Optionnel : compter d'abord le nombre de fichiers pour tqdm
    total_files = None
    if show_progress:
        total = 0
        for _dirpath, _dirnames, filenames in os.walk(root_path, followlinks=follow_symlinks):
            total += len(filenames)
        total_files = total

    walker = os.walk(root_path, followlinks=follow_symlinks)
    it = walker
    if show_progress:
        it = tqdm(walker, total=None, desc="Scanning folders", unit="folder")

    for dirpath, dirnames, filenames in it:
        for fname in filenames:
            try:
                fp = os.path.join(dirpath, fname)
                st = os.stat(fp)
                file_ext = Path(fname).suffix.lower().lstrip('.')
                item = {
                    "title": fname,
                    "extension": file_ext if file_ext != "" else None,
                    "path": fp,
                    "relative_path": safe_relpath(fp, root_path),
                    "folder": dirpath,
                    "size_bytes": st.st_size if hasattr(st, "st_size") else None,
                    "created": format_timestamp(getattr(st, "st_ctime", None)),
                    "modified": format_timestamp(getattr(st, "st_mtime", None)),
                }
                results.append(item)
            except Exception as e:
                # ne pas échouer sur un fichier inacessible
                print(f"Warning: couldn't stat file {os.path.join(dirpath, fname)}: {e}", file=sys.stderr)
    # tqdm wrapper above counts folders not files; create a more informative progress on files if large
    if show_progress:
        print(f"Scanned {len(results)} files under {root_path}")
    return results

def export_json(results, out_path):
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    print(f"JSON exporté -> {out_path}")

def export_excel(results, out_path):
    if pd is None:
        raise RuntimeError("pandas non installé. Installer : pip install pandas openpyxl")
    df = pd.DataFrame(results)
    df.to_excel(out_path, index=False)
    print(f"Excel exporté -> {out_path}")

def export_csv(results, out_path):
    if pd is None:
        raise RuntimeError("pandas non installé. Installer : pip install pandas")
    df = pd.DataFrame(results)
    df.to_csv(out_path, index=False)
    print(f"CSV exporté -> {out_path}")

def export_word(results, out_path, max_rows=2000):
    if Document is None:
        raise RuntimeError("python-docx non installé. Installer : pip install python-docx")
    doc = Document()
    doc.add_heading('Scan arborescence', level=1)
    doc.add_paragraph(f"Date: {datetime.now().isoformat(sep=' ', timespec='seconds')}")
    doc.add_paragraph(f"Nombre de fichiers: {len(results)}")

    # créer table (limiter le nombre de lignes pour éviter des .docx énormes)
    headers = ["title", "extension", "relative_path", "folder", "size_bytes", "modified"]
    rows = results[:max_rows]

    table = doc.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for r in rows:
        row_cells = table.add_row().cells
        for i, h in enumerate(headers):
            val = r.get(h, "")
            row_cells[i].text = "" if val is None else str(val)

    if len(results) > max_rows:
        doc.add_paragraph(f"... truncated: only first {max_rows} rows included in this .docx")

    doc.save(out_path)
    print(f"Word exporté -> {out_path}")

# ---------------------------
# CLI / Main
# ---------------------------
def choose_folder_dialog():
    # ouvre une boîte de dialogue simple (tkinter) si disponible
    try:
        import tkinter as tk
        from tkinter import filedialog
    except Exception:
        return None

    root = tk.Tk()
    root.withdraw()
    folder_selected = filedialog.askdirectory(title="Choisir le dossier racine à scanner")
    root.destroy()
    return folder_selected if folder_selected else None

def main():
    parser = argparse.ArgumentParser(description="Scanner une arborescence et exporter JSON/Excel/Word.")
    parser.add_argument("root", nargs="?", help="Chemin du dossier racine (si absent, boîte de dialogue sera ouverte)")
    parser.add_argument("--out-dir", "-o", help="Dossier de sortie pour les exports (défaut: dossier courant)", default=".")
    parser.add_argument("--no-word", action="store_true", help="Ne pas générer le .docx")
    parser.add_argument("--no-excel", action="store_true", help="Ne pas générer l'excel")
    parser.add_argument("--no-json", action="store_true", help="Ne pas générer json")
    parser.add_argument("--follow-symlinks", action="store_true", help="Suivre les liens symboliques")
    args = parser.parse_args()

    root = args.root
    if not root:
        root = choose_folder_dialog()
    if not root:
        print("Aucun dossier racine fourni. Fin.")
        sys.exit(1)
    if not os.path.isdir(root):
        print(f"Le chemin fourni n'est pas un dossier valide: {root}")
        sys.exit(1)

    print(f"Scanning root: {root}")
    results = scan_folder(root, follow_symlinks=args.follow_symlinks, show_progress=True)

    out_dir = os.path.abspath(args.out_dir)
    os.makedirs(out_dir, exist_ok=True)
    base_name = f"scan_{Path(root).name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    json_path = os.path.join(out_dir, base_name + ".json")
    excel_path = os.path.join(out_dir, base_name + ".xlsx")
    csv_path = os.path.join(out_dir, base_name + ".csv")
    word_path = os.path.join(out_dir, base_name + ".docx")

    if not args.no_json:
        export_json(results, json_path)
    if not args.no_excel:
        try:
            export_excel(results, excel_path)
        except Exception as e:
            print(f"Excel export failed: {e}. Trying CSV instead.")
            try:
                export_csv(results, csv_path)
            except Exception as e2:
                print(f"CSV export failed: {e2}")
    if not args.no_word:
        try:
            export_word(results, word_path)
        except Exception as e:
            print(f"Word export failed: {e}")

    print("Terminé.")

if __name__ == "__main__":
    main()
